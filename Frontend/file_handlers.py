# file_handlers.py
import os
import pdfplumber
from docx import Document as DocxDocument
from langchain.docstore.document import Document
import tempfile

def load_file(filepath):
    file_ext = os.path.splitext(filepath)[1].lower()
    if file_ext == ".pdf":
        text = load_pdf(filepath)
    elif file_ext == ".txt":
        text = load_txt(filepath)
    elif file_ext == ".docx":
        text = load_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")
    
    from text_processing import KoreanTextSplitter
    textsplitter = KoreanTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = textsplitter.split_text(text)
    docs = [Document(page_content=doc) for doc in docs]
    return docs

def load_pdf(filepath):
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
    return text

def load_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    return text

def load_docx(filepath):
    doc = DocxDocument(filepath)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def load_documents(uploaded_file):
    if uploaded_file is not None:
        if isinstance(uploaded_file, str):
            # 파일 경로가 문자열로 전달된 경우
            file_ext = os.path.splitext(uploaded_file)[1].lower()
            temp_file_path = uploaded_file
        else:
            # 파일 객체가 전달된 경우
            file_ext = os.path.splitext(uploaded_file.name)[1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_file_path = temp_file.name

        try:
            if file_ext == ".pdf":
                text = load_pdf(temp_file_path)
            elif file_ext == ".txt":
                text = load_txt(temp_file_path)
            elif file_ext == ".docx":
                text = load_docx(temp_file_path)
            elif file_ext == ".hwp":
                text = load_hwp(temp_file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

            from text_processing import KoreanTextSplitter
            textsplitter = KoreanTextSplitter(chunk_size=1000, chunk_overlap=200)
            docs = textsplitter.split_text(text)
            docs = [Document(page_content=doc) for doc in docs]
            return docs
        finally:
            if not isinstance(uploaded_file, str):
                os.unlink(temp_file_path)  # 임시 파일 삭제 (파일 객체로 전달된 경우에만)
    else:
        return None

import olefile
import zlib
import struct

def load_hwp(filename):
    f = olefile.OleFileIO(filename)
    dirs = f.listdir()

    # HWP 파일 검증
    if ["FileHeader"] not in dirs or \
       ["\x05HwpSummaryInformation"] not in dirs:
        raise Exception("Not Valid HWP.")

    # 문서 포맷 압축 여부 확인
    header = f.openstream("FileHeader")
    header_data = header.read()
    is_compressed = (header_data[36] & 1) == 1

    # Body Sections 불러오기
    nums = []
    for d in dirs:
        if d[0] == "BodyText":
            nums.append(int(d[1][len("Section"):]))
    sections = ["BodyText/Section"+str(x) for x in sorted(nums)]

    # 전체 text 추출
    text = ""
    for section in sections:
        bodytext = f.openstream(section)
        data = bodytext.read()
        if is_compressed:
            unpacked_data = zlib.decompress(data, -15)
        else:
            unpacked_data = data
    
        # 각 Section 내 text 추출    
        section_text = ""
        i = 0
        size = len(unpacked_data)
        while i < size:
            header = struct.unpack_from("<I", unpacked_data, i)[0]
            rec_type = header & 0x3ff
            rec_len = (header >> 20) & 0xfff

            if rec_type in [67]:
                rec_data = unpacked_data[i+4:i+4+rec_len]
                section_text += rec_data.decode('utf-16')
                section_text += "\n"

            i += 4 + rec_len

        text += section_text
        text += "\n"

    return text